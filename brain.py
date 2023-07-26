'''
this model based on Character-Level LSTM in PyTorch form udacity
recurrent-neural-networks/char-rnn
we also use word2vec-embeddings 
'''
#from .network import Model
#from .train import Train

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
from .network import TextRnn as Model
from .train import TexTrainer
import torch
from torch import nn
import torch.nn.functional as F
#from local library
from .utiliy import get_batches
#Load and prepare the data

import cv2
from imutils.video import VideoStream
from imutils.video import FPS
import argparse
import imutils
import pytesseract
from PIL import Image
import PIL.Image
import wand

TESSERACT_LOCATION=r'C:\Users\youssri.ahmed\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
class Thinking():    
    def __init__(self,folder,output):
        self.show_data
        self.folder=folder
        self.output=output
    def show_data(self,schema,table_name,*args,word=True,sentince=True):
        
        from libs.andlsy.analysis import Select,Data_db
        #from ..analysis import Select,Data_db
        #self.show_data.__init__(self, list(args))
        #os.chdir(self.folder)
        
        connection_type=Data_db()
        reader_item=connection_type.table_data(schema,table_name,*args)
        #https://towardsdatascience.com/multi-class-text-classification-with-lstm-1590bee1bd17
        #print("Info",reader_item.info())
        print("reader_item",reader_item[:10])
        col1=0
        col=list(args)
        col=col[col1]
        cols=col.split(",")
        
        col_name=cols[0]
        col_id=cols[1]


        counts_names=reader_item[col_name].value_counts()
        print("______names categories________",counts_names[5])

        counts_id=reader_item[col_id].value_counts()
        print("______id categories________",counts_id)
        print("reader_item col_name",col_name,reader_item[col_name])
        #reader_inside_category=reader_item.reader_inside_category.value_counts()
        # encode the text and map each character to an integer and vice versa

        if word:
            # we create two dictionaries:
                    #data pre-process for names
            chars_names = tuple(set(counts_names))
            int2char_name = dict(enumerate(counts_names))
            char2int_name = {ch: ii for ii, ch in int2char_name.items()}
            encoded_names= np.array([int2char_name[ch] for ch in counts_names])
            #========================================
                #first deictionaries for id
            chars_id = tuple(set(counts_id))
            int2char = dict(enumerate(chars_id))
            char2int = {ch: ii for ii, ch in int2char.items()}
            
            ## encode the text
            encoded_id = np.array([char2int[ch] for ch in counts_id])
            
        if sentince:
            #transfere dataframe to list
            # معايير تحديد الجمل الاسمية والفعلية            
            tow_words=reader_item.split("و")
            print("test___________ tow_words",tow_words.head(5))
            #data pre-process for names
            chars_names = tuple(set(counts_names))
            int2char_name = dict(enumerate(counts_names))
            char2int_name = {ch: ii for ii, ch in int2char_name.items()}
            encoded_names= np.array([int2char_name[ch] for ch in counts_names])
            #========================================
            # create id
                #first deictionaries for id
            chars_id = tuple(set(counts_id))
            int2char = dict(enumerate(chars_id))
            char2int = {ch: ii for ii, ch in int2char.items()}
            
            ## encode the text
            encoded_id = np.array([char2int[ch] for ch in counts_id])        
        print("encoded_id categories",encoded_id)

        print("____________ char2int_name ",char2int_name)
        ## encode the text

        print("encoded_name categories",encoded_names)

        def one_hot_encode(self,arr, n_labels):
            '''
            LSTM expects an input that is one-hot encoded_id meaning that each character 
            is converted into an integer (via our created dictionary) 
            and then converted into a column vector where only it's corresponding integer index will have the value of 1 and the rest of the vector will be filled with 0's. Since we're one-hot encoding the data
            '''
            # Initialize the the encoded_id array
            one_hot = np.zeros((np.multiply(*arr.shape), n_labels), dtype=np.float32)
            
            # Fill the appropriate elements with ones
            one_hot[np.arange(one_hot.shape[0]), arr.flatten()] = 1.
            
            # Finally reshape it to get back to the original array
            one_hot = one_hot.reshape((*arr.shape, n_labels))
            
            return one_hot
        # check that the function works as expected
        test_seq = np.array([[3, 5, 1]])
        one_hot = one_hot_encode(self,test_seq, 8)

        print("____one_hot____",one_hot)
        
        self.batches = get_batches(encoded_id, 8, 50)
        print("____fast test_____ counts_id.self.batches()",self.batches)
        self.x, self.y = next(self.batches)
        # printing out the first 10 items in a sequence
        print('x\n', self.x[:10, :10])
        print('\ny\n', self.y[:10, :10])

        self.batches_name = get_batches(encoded_col1, 8, 50)
        self.z, self.t = next(self.batches_name)
        # printing out the first 10 items in a sequence
        print('z\n', self.z[:10, :10])
        print('\ny\n', self.t[:10, :10])
        # check if GPU is available
        train_on_gpu = torch.cuda.is_available()
        if(train_on_gpu):
            print('Training on GPU!')
        else: 
            print('No GPU available, training on CPU; consider making n_epochs very small.')
        ## : set you model hyperparameters
        # define and print the net
        n_hidden=512
        n_layers=2

        #net = Model(chars_col1, n_hidden, n_layers)
        net = Model(chars_id, n_hidden, n_layers)
        print(net)
        batch_size = 128
        seq_length = 100
        n_epochs =  20 # start small if you are just testing initial behavior

        # train the model
        #TexTrainer(net, encoded_col1, epochs=n_epochs, batch_size=batch_size, seq_length=seq_length, lr=0.001, print_every=10)
        TexTrainer.train(net, encoded_id, epochs=n_epochs, batch_size=batch_size, seq_length=seq_length, lr=0.001, print_every=10)

        ## change the name, for saving multiple files

        model_name = 'rnn_x_epoch.net'

        checkpoint = {'n_hidden': net.n_hidden,
                    'n_layers': net.n_layers,
                    'state_dict': net.state_dict(),
                    'tokens': net.chars}

        with open(model_name, 'wb') as f:
            torch.save(checkpoint, f)
        # change the name, for saving multiple files
        
        model_name = 'rnn_x_epoch.net'

        checkpoint = {'n_hidden': net.n_hidden,
                    'n_layers': net.n_layers,
                    'state_dict': net.state_dict(),
                    'tokens': net.chars}

        with open(model_name, 'wb') as f:
            torch.save(checkpoint, f)
    def tesseract_ocr(self):
        import pytesseract
        from PIL import Image
        print('start reading letters')
        from pytesseract import image_to_string

        # img = Image.open('E:\Python projects\FaceDetect-master - Copy\image.jpg')
        img = Image.open(self.folder)

        pytesseract.pytesseract.tesseract_cmd = TESSERACT_LOCATION

        # text = image_to_string(img, lang="eng")
        text = image_to_string(img, lang="ara")

        print(text)
    def face_detect(self):
        import cv2
        import sys

        import time
        # Get user supplied values
        imagePath = sys.argv[1]
        cascPath = "haarcascade_frontalface_default.xml"

        # Create the haar cascade
        faceCascade = cv2.CascadeClassifier(cascPath)

        # Read the image
        image = cv2.imread(imagePath)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Detect faces in the image
        faces = faceCascade.detectMultiScale(
            gray,
            scaleFactor=1.1,
            minNeighbors=5,
            minSize=(30, 30)
            #flags = cv2.CV_HAAR_SCALE_IMAGE
        )

        print("Found {0} faces!".format(len(faces)))

        # Draw a rectangle around the faces
        for (x, y, w, h) in faces:
            cv2.rectangle(image, (x, y), (x+w, y+h), (0, 255, 0), 2)

        cv2.imshow("Faces found", image)
        cv2.waitKey(0)
    def live(self):

        cap = cv2.VideoCapture(0)

        template = cv2.imread("E:/Python projects/MatchingImages/ahmed_after_cuttting.png",
                            cv2.IMREAD_GRAYSCALE)
        template1 = cv2.imread("E:/Python projects/MatchingImages/kamal2.png",
                            cv2.IMREAD_GRAYSCALE)
        template2 = cv2.imread("E:/Python projects/MatchingImages/medhat.png",
                            cv2.IMREAD_GRAYSCALE)

        w, h = template.shape[::-1]
        w1, h1 = template1.shape[::-1]
        w2, h2 = template2.shape[::-1]

        font = cv2.FONT_HERSHEY_PLAIN

        # Create the haar cascade
        faceCascade = cv2.CascadeClassifier("haarcascade_frontalface_default.xml")

        while (True):
            # Capture frame-by-frame
            # ret, frame = cap.read()
            _, frame = cap.read()
            # Our operations on the frame come here
            gray_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

            res = cv2.matchTemplate(gray_frame, template, cv2.TM_CCOEFF_NORMED)
            res1 = cv2.matchTemplate(gray_frame, template1, cv2.TM_CCOEFF_NORMED)
            res2 = cv2.matchTemplate(gray_frame, template2, cv2.TM_CCOEFF_NORMED)

            thres = 1
            loc = np.where(res >= thres)

            loc1 = np.where(res1 >= thres)
            loc2 = np.where(res2 >= thres)
            # Detect faces in the image
            faces = faceCascade.detectMultiScale(
                gray_frame,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(30, 30)
                # flags = cv2.CV_HAAR_SCALE_IMAGE
            )

            print("Found {0} faces!".format(len(faces)))

            # Draw a rectangle around the faces
            for (x, y, w, h) in faces:
                cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 2)
                for pt in zip(*loc[::-1]):
                    # cv2.rectangle(frame, pt, (pt[0] + w, pt[1] + h), (0, 255, 0), 3)
                    cv2.putText(frame, 'Ahmed', (pt[0] + w, pt[1] + h), font, 3, (255, 255, 255), 1)
                for pt1 in zip(*loc1[::-1]):
                    # cv2.rectangle(frame, pt1, (pt1[0] + w1, pt1[1] + h1), (0, 255, 0), 3)
                    cv2.putText(frame, 'Kamal', (pt1[0] + w1, pt1[1] + h1), font, 3, (255, 255, 255), 1)
                for pt2 in zip(*loc2[::-1]):
                    # cv2.rectangle(frame, pt2, (pt2[0] + w2, pt2[1] + h2), (0, 255, 0), 3)
                    cv2.putText(frame, 'Medhat', (pt2[0] + w2, pt2[1] + h2), font, 3, (255, 255, 255), 1)

            # Display the resulting frame
            cv2.imshow('frame', frame)
            if cv2.waitKey(1) & 0xFF == ord('q'):
                break

        # When everything done, release the capture
        cap.release()
        cv2.destroyAllWindows()
    def text_recogition(self):
        # from wand.image import Image as wi

        #
        # pdf = wi(filename="E:\Python projects\FaceDetect-master - Copy\AhmedRoid.pdf, resolution=300")
        # pdfimage = pdf.convert("jpeg")
        # i=1
        # for img in pdfimage.sequence:
        #     page = wi(image=img)
        #     page.save(filename="E:\Python projects\FaceDetect-master - Copy\Images" + "\'" + str(i) + ".jpg")
        #     i +=1


        import os
        from pdf2image import convert_from_path
        from pytesseract import image_to_string
        from docx import Document
        from docx.shared import Inches

        document = Document()
        document.add_heading('Document Title', 0)

        pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf2"
        pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf3"
        pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf1"

        os.chdir(pdf_dir)
        for pdf_file in os.listdir(pdf_dir):
            if pdf_file.endswith(".pdf"):
                pages = convert_from_path(pdf_file, 250)
                # pdf_file = pdf_file[:-4]
                pdf_file = pdf_file[0]
                for page in pages:
                    page.save("%s-page%d.jpg" % (pdf_file, pages.index(page)), "JPEG")
                    # page.p.add_run('bold').bold = Truesave("E:\Python projects\FaceDetect-master - Copy\Images-page%d.jpg" % (pages.index(page)), "JPEG")
                for image_file in os.listdir(pdf_dir):
                    if image_file.endswith(".jpg"):
                        # print('gggg')
                        img = Image.open(image_file)
                        text = image_to_string(img, lang="ara")
                        #print(text)
                        #document.add_paragraph("Ahmed mohamed", style='List Number')
                        document.add_paragraph(text, style='List Number')
                        #p.add_run('bold').bold = True
                        document.add_page_break()



        document.save('demo.docx')

        # for image_file in os.listdir(pdf_dir):
        #
        #     if image_file.endswith(".jpg"):
        #         # print('gggg')
        #         img = Image.open(image_file)
        #         text = image_to_string(img, lang="ara")
        #         print(text)
