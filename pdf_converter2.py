import cv2
import numpy as np
from imutils.video import VideoStream
from imutils.video import FPS
import argparse
import imutils
import pytesseract
from PIL import Image
import PIL.Image
import cv2
import numpy as np
import pytesseract
from PIL import Image
import wand
# from wand.image import Image as wi

#
# pdf = wi(filename="E:\Python projects\FaceDetect-master - Copy\AhmedRoid.pdf, resolution=300")
# pdfimage = pdf.convert("jpeg")
# i=1
# for img in pdfimage.sequence:
#     page = wi(image=img)
#     page.save(filename="E:\Python projects\FaceDetect-master - Copy\Images" + "\'" + str(i) + ".jpg")
#     i +=1


import os
from pdf2image import convert_from_path
from pytesseract import image_to_string
from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('Document Title', 0)

pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf2"
pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf3"
pdf_dir = r"E:\Python projects\FaceDetect-master - Copy\pdf1"

os.chdir(pdf_dir)
for pdf_file in os.listdir(pdf_dir):
    if pdf_file.endswith(".pdf"):
        pages = convert_from_path(pdf_file, 250)
        # pdf_file = pdf_file[:-4]
        pdf_file = pdf_file[0]
        for page in pages:
            page.save("%s-page%d.jpg" % (pdf_file, pages.index(page)), "JPEG")
            # page.p.add_run('bold').bold = Truesave("E:\Python projects\FaceDetect-master - Copy\Images-page%d.jpg" % (pages.index(page)), "JPEG")
        for image_file in os.listdir(pdf_dir):
            if image_file.endswith(".jpg"):
                # print('gggg')
                img = Image.open(image_file)
                text = image_to_string(img, lang="ara")
                #print(text)
                #document.add_paragraph("Ahmed mohamed", style='List Number')
                document.add_paragraph(text, style='List Number')
                #p.add_run('bold').bold = True
                document.add_page_break()



document.save('demo.docx')

# for image_file in os.listdir(pdf_dir):
#
#     if image_file.endswith(".jpg"):
#         # print('gggg')
#         img = Image.open(image_file)
#         text = image_to_string(img, lang="ara")
#         print(text)
